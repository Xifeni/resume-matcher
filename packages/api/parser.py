import os
import shutil
import subprocess
import tempfile

import pdfplumber
from docx import Document

ALLOWED_UPLOAD_EXTENSIONS = frozenset({".pdf", ".doc", ".docx"})


def extension_from_filename(filename: str) -> str:
    if not filename:
        raise ValueError("Пустое имя файла")
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_UPLOAD_EXTENSIONS:
        raise ValueError(
            f"Неподдерживаемый формат {ext or '(нет)'}; допустимо: "
            + ", ".join(sorted(ALLOWED_UPLOAD_EXTENSIONS))
        )
    return ext


def extract_text_from_pdf(pdf_path: str) -> str:
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
    return text


def extract_text_from_docx(docx_path: str) -> str:
    doc = Document(docx_path)
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def extract_text_from_doc(doc_path: str) -> str:
    exe = shutil.which("antiword")
    if not exe:
        raise ValueError(
            "Для .doc нужна утилита antiword в системе (в Docker-образе API она установлена)."
        )
    r = subprocess.run(
        [exe, doc_path],
        capture_output=True,
        text=True,
        timeout=120,
    )
    if r.returncode != 0:
        err = (r.stderr or r.stdout or "").strip() or f"код {r.returncode}"
        raise ValueError(f"antiword: {err}")
    return (r.stdout or "").strip()


def extract_text_from_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    if ext == ".docx":
        return extract_text_from_docx(path)
    if ext == ".doc":
        return extract_text_from_doc(path)
    raise ValueError(f"Неизвестное расширение файла: {ext}")


def extract_text_from_upload(file_storage) -> str:
    """Сохраняет загрузку во временный файл с корректным суффиксом, извлекает текст, удаляет файл."""
    if not file_storage or not file_storage.filename:
        raise ValueError("Файл не передан")
    ext = extension_from_filename(file_storage.filename)
    fd, path = tempfile.mkstemp(suffix=ext)
    os.close(fd)
    try:
        file_storage.save(path)
        return extract_text_from_file(path)
    finally:
        if os.path.isfile(path):
            os.remove(path)


def save_upload_to_temp_path(file_storage) -> str:
    """Сохраняет загрузку во временный файл; вызывающий обязан удалить путь после extract_text_from_file."""
    if not file_storage or not file_storage.filename:
        raise ValueError("Файл не передан")
    ext = extension_from_filename(file_storage.filename)
    fd, path = tempfile.mkstemp(suffix=ext)
    os.close(fd)
    file_storage.save(path)
    return path
